import pandas as pd
import io
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
import datetime

def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode('utf-8')


def export_pdf(df: pd.DataFrame, title="Data Export") -> bytes:
    # A4 landscape orientation for better table fit
    pdf = FPDF(orientation="landscape", unit="mm", format="A4")
    pdf.add_page()
    
    # Aesthetic Title
    pdf.set_font("helvetica", "B", 16)
    pdf.set_text_color(41, 128, 185) # Flat Blue
    pdf.cell(0, 10, title, ln=True, align="C")
    
    # Subtitle with timestamp
    pdf.set_text_color(100, 100, 100) # Gray
    pdf.set_font("helvetica", "I", 10)
    pdf.cell(0, 8, f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="C")
    pdf.ln(5)
    
    # Limit rows for PDF performance and to avoid overflowing memory
    df_preview = df.head(1000)
    
    # Smart width calculations prioritizing available 277mm
    usable_width = 277
    cols = df_preview.columns
    col_widths = []
    for col in cols:
        content_max = df_preview[col].astype(str).map(len).max() if not df_preview.empty else 5
        header_max = len(str(col))
        max_len = max(content_max, header_max)
        col_widths.append(min(max_len * 2 + 5, (usable_width / len(cols)) + 20))
    
    total_w = sum(col_widths)
    if total_w > usable_width:
        col_widths = [w * (usable_width / total_w) for w in col_widths]
        
    # Table Header Design
    pdf.set_fill_color(41, 128, 185) 
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 8)
    
    for w, col in zip(col_widths, cols):
        pdf.cell(w, 8, str(col)[:20], border=1, fill=True, align="C")
    pdf.ln()
    
    # Table Data Design
    pdf.set_text_color(50, 50, 50)
    pdf.set_font("helvetica", size=7)
    
    fill = False
    for _, row in df_preview.iterrows():
        # Soft striping for readability
        if fill:
            pdf.set_fill_color(245, 247, 250)
        else:
            pdf.set_fill_color(255, 255, 255)
            
        for w, item in zip(col_widths, row):
            val_str = str(item).replace('\n', ' ')
            max_char = max(3, int(w/1.5) - 2)
            if len(val_str) > max_char:
                val_str = val_str[:max_char] + ".."
            pdf.cell(w, 6, val_str, border=1, fill=True)
        pdf.ln()
        fill = not fill
        
    return bytes(pdf.output())

def export_docx(df: pd.DataFrame, title="Data Export") -> bytes:
    doc = Document()
    
    # Colorful Word Title
    heading = doc.add_heading(title, 0)
    
    # Subtitle
    p = doc.add_paragraph(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    df_preview = df.head(1000)
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # Professional Word Theme application
    table.style = 'Light Shading Accent 1'
    
    # Build Headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    # Build Data Rows
    for _, row in df_preview.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            val_str = str(item)
            if len(val_str) > 1000:
                val_str = val_str[:997] + "..."
            row_cells[i].text = val_str
            
    # Serialize to memory
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
